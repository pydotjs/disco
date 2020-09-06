# import packages
import docx

from disco.file.File import File
import disco.logs
logger = disco.logs.get(__name__)

class DocxFile(File):

    @staticmethod
    def open_file(path):
        # open the docx file
        return docx.Document(str(path.resolve()))

    def inspect_chunks(self):
        file = self.file
        return len(file.paragraphs)

    def yield_chunks(self):
        file = self.file
        for para in file.paragraphs:
            yield para.text

